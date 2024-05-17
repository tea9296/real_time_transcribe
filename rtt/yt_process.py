import subprocess
#from pytube import YouTube
from pytubefix import YouTube
from pathlib import Path
import whisper
import datetime
import shutil
from docx import Document


# Function to download a YouTube video
def download_video(url, output_dir):
    yt = YouTube(url)
    stream = yt.streams.filter(only_audio=True).first()
    download_path = Path(output_dir)
    download_path.mkdir(
        parents=True,
        exist_ok=True)  # Create the directory if it doesn't exist
    # use datetime to name the file
    title = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    file_name = stream.title
    output_file = download_path / (title + ".mp4")
    stream.download(output_path=download_path, filename=title + ".mp4")
    return output_file, file_name


# Function to convert video to WAV using FFmpeg
def convert_to_wav(input_file, output_file):
    subprocess.run([
        "ffmpeg", "-i", input_file, "-acodec", "pcm_s16le", "-ar", "16000",
        output_file
    ])


# Function to filter vocals using Demucs
def filter_vocals(input_file):
    subprocess.run(["demucs", "--two-stems=vocals", input_file])


# Function to transcribe audio using OpenAI Whisper
def transcribe_audio(input_file,
                     model_name: str,
                     language: str = "English",
                     timestamps=False):

    model = whisper.load_model(set_model_name(model_name))
    if language == "Chinese" or language == "zh":
        result = model.transcribe(input_file.__str__(),
                                  initial_prompt="以下是普通話的句子。",
                                  verbose=True,
                                  word_timestamps=True)
    else:
        result = model.transcribe(input_file.__str__(),
                                  verbose=True,
                                  word_timestamps=True)

    res = ""

    ## return with timestamps##
    if timestamps:

        for segment in result['segments']:
            start, end, text = segment["start"], segment["end"], segment[
                "text"]
            line = f"[{format_timestamp(start)} --> {format_timestamp(end)}] {text}"
            res += line + "\n"

        return res

    for segment in result['segments']:
        res += segment['text'] + "\n"

    return res


def save_doc(text, output_file_name):
    # Create a new document
    doc = Document()

    # Add a title to the document
    doc.add_heading(output_file_name.split('.')[:-1], 0)

    # Your transcription content, including both Chinese and English
    # transcription = text.replace('。 ', '。\n\n').replace('! ', '!\n\n').replace(
    #     '? ', '?\n\n').replace('. ', '.\n\n')

    # Add the transcription content to the document
    doc.add_paragraph(text)

    # Save the document to a file
    doc.save(output_file_name)

    return


# Main function to perform the entire process
def process_youtube_video(url: str,
                          output_file_name: str = "",
                          model_name: str = "medium",
                          language: str = "English"):

    output_dir = ".temp"

    video_file, file_name = download_video(url, output_dir)
    wav_file = Path(output_dir) / (video_file.stem + ".wav")

    if output_file_name == "":
        output_file_name = "./" + file_name.replace('/', '_').replace(
            ' ', '').replace('\"', '').replace('\'', '') + ".docx"

    if output_file_name.split('.')[-1] != 'docx':
        timestamps = True
    else:
        timestamps = False

    separated_audio = f"./separated/htdemucs/{video_file.stem}/vocals.wav"

    convert_to_wav(video_file, wav_file)
    filter_vocals(wav_file)
    transcription = transcribe_audio(separated_audio, model_name, language,
                                     timestamps)

    # Delete temporary files
    video_file.unlink()
    wav_file.unlink()
    #separated_audio.unlink()
    try:
        shutil.rmtree('./separated/')
        shutil.rmtree('./.temp/')
    except OSError as e:
        print(f"Error: {e}")

    if output_file_name.split('.')[-1] == 'docx':
        save_doc(transcription, output_file_name)
    else:
        with open(output_file_name, 'w', encoding="utf-8") as f:
            f.write(transcription)

    return transcription


# Main function to perform the entire process
def process_wav_audio(input_file: str,
                      output_file_name: str = "",
                      model_name: str = "medium",
                      language: str = "English"):

    wav_file = Path(input_file)
    if output_file_name == "":
        output_file_name = "transcription.docx"

    if output_file_name.split('.')[-1] != 'docx':
        timestamps = True
    else:
        timestamps = False

    separated_audio = f"./separated/htdemucs/{wav_file.stem}/vocals.wav"

    filter_vocals(wav_file)
    transcription = transcribe_audio(separated_audio, model_name, language,
                                     timestamps)

    #separated_audio.unlink()
    try:
        shutil.rmtree('./separated/')
    except OSError as e:
        print(f"Error: {e}")

    if output_file_name.split('.')[-1] == 'docx':
        save_doc(transcription, output_file_name)
    else:
        with open(output_file_name, 'w', encoding="utf-8") as f:
            f.write(transcription)

    return transcription


def set_model_name(model_name: str):

    all_model = whisper.available_models()
    if model_name in all_model:
        return model_name
    return "medium"


def format_timestamp(seconds: float,
                     always_include_hours: bool = False,
                     decimal_marker: str = "."):
    assert seconds >= 0, "non-negative timestamp expected"
    milliseconds = round(seconds * 1000.0)

    hours = milliseconds // 3_600_000
    milliseconds -= hours * 3_600_000

    minutes = milliseconds // 60_000
    milliseconds -= minutes * 60_000

    seconds = milliseconds // 1_000
    milliseconds -= seconds * 1_000

    hours_marker = f"{hours:02d}:" if always_include_hours or hours > 0 else ""
    return (
        f"{hours_marker}{minutes:02d}:{seconds:02d}{decimal_marker}{milliseconds:03d}"
    )
